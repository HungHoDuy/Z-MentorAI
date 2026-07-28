import datetime
import io
import json
import re
import zipfile
from dataclasses import dataclass
from typing import Any

from docx import Document
from fastapi import HTTPException
from google.api_core.client_options import ClientOptions
from google.cloud import documentai
from google.cloud import storage
from pypdf import PdfReader
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.config import logger, settings
from cv_intake.repository import update_cv_document
from cv_extraction.schemas import CvExtractionResult


MIN_EXTRACTED_TEXT_CHARS = 80
EXTRACTION_VERSION = "cv-extraction-v3"


@dataclass(frozen=True)
class TextQuality:
    usable: bool
    char_count: int
    word_count: int
    alphanumeric_ratio: float
    words_per_page: float
    sparse_page_count: int
    short_line_ratio: float
    duplicate_line_ratio: float
    mojibake_ratio: float
    reasons: list[str]


def assess_pdf_text_quality(
    text: str,
    page_count: int | None,
    page_char_counts: list[int] | None = None,
) -> TextQuality:
    stripped = (text or "").strip()
    non_space = re.sub(r"\s", "", stripped)
    alphanumeric = sum(char.isalnum() for char in non_space)
    words = re.findall(r"\b\w{2,}\b", stripped, flags=re.UNICODE)
    pages = max(page_count or 1, 1)
    ratio = alphanumeric / max(len(non_space), 1)
    words_per_page = len(words) / pages
    sparse_page_count = sum(1 for count in page_char_counts or [] if count < 40)
    lines = [" ".join(line.split()) for line in stripped.splitlines() if line.strip()]
    short_lines = [line for line in lines if len(line.split()) <= 2]
    normalized_lines = [line.casefold() for line in lines if len(line) >= 8]
    duplicate_line_count = max(0, len(normalized_lines) - len(set(normalized_lines)))
    short_line_ratio = len(short_lines) / max(len(lines), 1)
    duplicate_line_ratio = duplicate_line_count / max(len(normalized_lines), 1)
    mojibake_markers = set("ÃÂâãæçåð�")
    mojibake_ratio = sum(char in mojibake_markers for char in stripped) / max(len(stripped), 1)
    reasons = []
    if len(stripped) < MIN_EXTRACTED_TEXT_CHARS:
        reasons.append("too_few_characters")
    if len(words) < max(15, pages * 12):
        reasons.append("too_few_words")
    if words_per_page < 12:
        reasons.append("low_words_per_page")
    if ratio < 0.65:
        reasons.append("low_alphanumeric_ratio")
    if sparse_page_count:
        reasons.append("sparse_pdf_pages")
    if len(lines) >= 20 and short_line_ratio > 0.85:
        reasons.append("fragmented_reading_order")
    if len(normalized_lines) >= 20 and duplicate_line_ratio > 0.65:
        reasons.append("duplicated_layout_text")
    if mojibake_ratio > 0.02:
        reasons.append("encoding_mojibake")
    return TextQuality(
        usable=not reasons,
        char_count=len(stripped),
        word_count=len(words),
        alphanumeric_ratio=round(ratio, 3),
        words_per_page=round(words_per_page, 2),
        sparse_page_count=sparse_page_count,
        short_line_ratio=round(short_line_ratio, 3),
        duplicate_line_ratio=round(duplicate_line_ratio, 3),
        mojibake_ratio=round(mojibake_ratio, 3),
        reasons=reasons,
    )


def utc_now() -> str:
    return datetime.datetime.now(datetime.timezone.utc).isoformat()


def log_extraction_event(event: str, **fields: Any) -> None:
    logger.info(json.dumps({"event": event, **fields}, ensure_ascii=True, default=str))


def extract_pdf_text(content: bytes) -> tuple[str, int, list[int]]:
    reader = PdfReader(io.BytesIO(content))
    page_texts = []
    for page in reader.pages:
        try:
            page_texts.append(page.extract_text(extraction_mode="layout") or "")
        except (TypeError, ValueError):
            page_texts.append(page.extract_text() or "")
    text = "\n\n".join(part.strip() for part in page_texts if part.strip()).strip()
    return text, len(reader.pages), [len(part.strip()) for part in page_texts]


def get_document_ai_processor_name() -> str | None:
    if settings.document_ai_processor_name:
        return settings.document_ai_processor_name
    if not settings.document_ai_processor_id:
        return None

    client = documentai.DocumentProcessorServiceClient(
        client_options=ClientOptions(
            api_endpoint=f"{settings.document_ai_location}-documentai.googleapis.com",
        )
    )
    return client.processor_path(
        settings.google_cloud_project,
        settings.document_ai_location,
        settings.document_ai_processor_id,
    )


def extract_text_with_document_ai(content: bytes, mime_type: str) -> str:
    processor_name = get_document_ai_processor_name()
    if not processor_name:
        raise RuntimeError("Document AI processor is not configured.")

    client = documentai.DocumentProcessorServiceClient(
        client_options=ClientOptions(
            api_endpoint=f"{settings.document_ai_location}-documentai.googleapis.com",
        )
    )
    request = documentai.ProcessRequest(
        name=processor_name,
        raw_document=documentai.RawDocument(
            content=content,
            mime_type=mime_type,
        ),
    )
    result = client.process_document(request=request)
    return (result.document.text or "").strip()


def extract_pdf_text_with_document_ai(content: bytes) -> str:
    return extract_text_with_document_ai(content, "application/pdf")


def extract_docx_images(content: bytes, limit: int = 10) -> list[tuple[bytes, str]]:
    mime_types = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
    }
    images = []
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        for member in archive.infolist():
            lowered = member.filename.casefold()
            extension = next(
                (suffix for suffix in mime_types if lowered.endswith(suffix)),
                None,
            )
            if not extension or not lowered.startswith("word/media/"):
                continue
            images.append((archive.read(member), mime_types[extension]))
            if len(images) >= limit:
                break
    return images


def extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    blocks = []

    def add_text(value: str) -> None:
        text = " ".join((value or "").split())
        if text and text not in blocks:
            blocks.append(text)

    for block in document.iter_inner_content():
        if isinstance(block, Paragraph):
            add_text(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
                if cells:
                    add_text(" | ".join(cells))

    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                add_text(paragraph.text)
            for table in part.tables:
                for row in table.rows:
                    cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
                    if cells:
                        add_text(" | ".join(cells))

    for text_box in document.element.xpath(".//w:txbxContent"):
        add_text(" ".join(node.text or "" for node in text_box.xpath(".//w:t")))

    return "\n".join(blocks).strip()


def build_artifact_object_name(original_object: str, filename: str) -> str:
    prefix = original_object.rsplit("/", 1)[0]
    return f"{prefix}/{filename}"


def upload_text_artifact(bucket: Any, object_name: str, content: str, content_type: str) -> str:
    blob = bucket.blob(object_name)
    blob.upload_from_string(content.encode("utf-8"), content_type=content_type)
    return f"gs://{bucket.name}/{object_name}"


async def extract_cv_text(document: dict) -> CvExtractionResult:
    cv_document_id = document["cv_document_id"]
    file_kind = document.get("file_kind")
    bucket_name = document.get("storage_bucket")
    object_name = document.get("storage_object")
    log_extraction_event(
        "cv_text_extraction_started",
        cv_document_id=cv_document_id,
        file_kind=file_kind,
        file_size_bytes=document.get("file_size_bytes"),
    )

    if file_kind not in {"pdf", "docx"}:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX CV extraction is supported.")
    if not bucket_name or not object_name:
        raise HTTPException(status_code=400, detail="CV document storage metadata is incomplete.")

    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(object_name)

    try:
        content = blob.download_as_bytes()
    except Exception as exc:
        logger.exception(
            "Failed to download CV from GCS",
            extra={
                "cv_document_id": cv_document_id,
                "bucket": bucket_name,
                "object_name": object_name,
                "error_type": type(exc).__name__,
            },
        )
        raise HTTPException(status_code=503, detail="Unable to download CV from GCS.") from exc

    page_count = None
    page_char_counts = None
    parser_type = "pypdf" if file_kind == "pdf" else "python_docx"
    extraction_quality = None
    try:
        if file_kind == "pdf":
            try:
                parsed_text, page_count, page_char_counts = extract_pdf_text(content)
            except Exception as parser_exc:
                if not get_document_ai_processor_name():
                    raise
                logger.warning(
                    "PyPDF extraction failed; using Document AI OCR",
                    extra={
                        "cv_document_id": cv_document_id,
                        "error_type": type(parser_exc).__name__,
                    },
                )
                parser_type = "document_ai_ocr"
                log_extraction_event(
                    "cv_document_ai_ocr_started",
                    cv_document_id=cv_document_id,
                    reason="pypdf_error",
                )
                parsed_text = extract_pdf_text_with_document_ai(content)
                log_extraction_event(
                    "cv_document_ai_ocr_completed",
                    cv_document_id=cv_document_id,
                    text_char_count=len(parsed_text),
                )
            else:
                extraction_quality = assess_pdf_text_quality(parsed_text, page_count, page_char_counts)
                logger.info(
                    "Evaluated PDF text quality",
                    extra={
                        "cv_document_id": cv_document_id,
                        "usable": extraction_quality.usable,
                        "char_count": extraction_quality.char_count,
                        "word_count": extraction_quality.word_count,
                        "words_per_page": extraction_quality.words_per_page,
                        "alphanumeric_ratio": extraction_quality.alphanumeric_ratio,
                        "mojibake_ratio": extraction_quality.mojibake_ratio,
                        "reasons": extraction_quality.reasons,
                    },
                )
                if not extraction_quality.usable:
                    if get_document_ai_processor_name():
                        parser_type = "document_ai_ocr"
                        log_extraction_event(
                            "cv_document_ai_ocr_started",
                            cv_document_id=cv_document_id,
                            reason="embedded_text_quality",
                            quality_reasons=extraction_quality.reasons,
                        )
                        parsed_text = extract_pdf_text_with_document_ai(content)
                        log_extraction_event(
                            "cv_document_ai_ocr_completed",
                            cv_document_id=cv_document_id,
                            text_char_count=len(parsed_text),
                        )
        else:
            parsed_text = extract_docx_text(content)
            embedded_images = extract_docx_images(content)
            if (
                embedded_images
                and len(parsed_text) < 500
                and get_document_ai_processor_name()
            ):
                ocr_parts = []
                for image_content, image_mime_type in embedded_images:
                    image_text = extract_text_with_document_ai(image_content, image_mime_type)
                    if image_text:
                        ocr_parts.append(image_text)
                if ocr_parts:
                    parsed_text = "\n\n".join([parsed_text, *ocr_parts]).strip()
                    parser_type = "python_docx+document_ai_ocr"
    except Exception as exc:
        await update_cv_document(cv_document_id, {
            "extraction_status": "failed",
            "extraction_error": str(exc),
            "extracted_at": utc_now(),
        })
        logger.exception(
            "Failed to extract CV text",
            extra={
                "cv_document_id": cv_document_id,
                "file_kind": file_kind,
                "parser_type": parser_type,
                "error_type": type(exc).__name__,
            },
        )
        raise HTTPException(status_code=422, detail="Unable to extract text from this CV.") from exc

    final_quality = (
        assess_pdf_text_quality(
            parsed_text,
            page_count,
            None if parser_type == "document_ai_ocr" else page_char_counts,
        )
        if file_kind == "pdf"
        else None
    )
    if len(parsed_text) < MIN_EXTRACTED_TEXT_CHARS or (final_quality and not final_quality.usable):
        await update_cv_document(cv_document_id, {
            "extraction_status": "needs_ocr",
            "parser_type": parser_type,
            "text_char_count": len(parsed_text),
            "page_count": page_count,
            "extracted_at": utc_now(),
            "extraction_error": "Extracted text is too short and OCR fallback is not configured or returned too little text.",
            "extraction_quality": final_quality.__dict__ if final_quality else None,
        })
        raise HTTPException(
            status_code=422,
            detail="Extracted text is too short. OCR fallback is not configured or returned too little text.",
        )

    extracted_at = utc_now()
    parsed_text_object = build_artifact_object_name(object_name, "parsed_text.txt")
    parsed_result_object = build_artifact_object_name(object_name, "parsed_result.json")
    parsed_result = {
        "cv_document_id": cv_document_id,
        "extraction_version": EXTRACTION_VERSION,
        "parser_type": parser_type,
        "ocr_fallback_used": "document_ai_ocr" in parser_type,
        "text_char_count": len(parsed_text),
        "page_count": page_count,
        "extraction_quality": final_quality.__dict__ if final_quality else None,
        "extracted_at": extracted_at,
    }

    try:
        parsed_text_uri = upload_text_artifact(
            bucket,
            parsed_text_object,
            parsed_text,
            "text/plain; charset=utf-8",
        )
        parsed_result_uri = upload_text_artifact(
            bucket,
            parsed_result_object,
            json.dumps(parsed_result, ensure_ascii=False, indent=2),
            "application/json; charset=utf-8",
        )
    except Exception as exc:
        logger.exception(
            "Failed to upload parsed CV artifacts",
            extra={
                "cv_document_id": cv_document_id,
                "bucket": bucket_name,
                "error_type": type(exc).__name__,
            },
        )
        raise HTTPException(status_code=503, detail="Unable to store parsed CV artifacts.") from exc

    await update_cv_document(cv_document_id, {
        "extraction_status": "completed",
        "extraction_version": EXTRACTION_VERSION,
        "parser_type": parser_type,
        "ocr_fallback_used": "document_ai_ocr" in parser_type,
        "text_char_count": len(parsed_text),
        "page_count": page_count,
        "extraction_quality": final_quality.__dict__ if final_quality else None,
        "parsed_text_gcs_uri": parsed_text_uri,
        "parsed_result_gcs_uri": parsed_result_uri,
        "parsed_text_object": parsed_text_object,
        "parsed_result_object": parsed_result_object,
        "extracted_at": extracted_at,
        "next_status": "pending_profile_extraction",
        "extraction_error": None,
    })
    log_extraction_event(
        "cv_text_extraction_completed",
        cv_document_id=cv_document_id,
        extraction_version=EXTRACTION_VERSION,
        parser_type=parser_type,
        ocr_fallback_used="document_ai_ocr" in parser_type,
        text_char_count=len(parsed_text),
        page_count=page_count,
        quality=final_quality.__dict__ if final_quality else None,
    )

    return CvExtractionResult(
        cv_document_id=cv_document_id,
        extraction_version=EXTRACTION_VERSION,
        parser_type=parser_type,
        ocr_fallback_used="document_ai_ocr" in parser_type,
        text_char_count=len(parsed_text),
        page_count=page_count,
        parsed_text_gcs_uri=parsed_text_uri,
        parsed_result_gcs_uri=parsed_result_uri,
        extracted_at=extracted_at,
        message_vi=(
            "Profile Scanner đã trích xuất nội dung CV thành công. "
            "Bước tiếp theo là chuẩn hóa kỹ năng, học vấn, kinh nghiệm và đánh giá benchmark theo target role."
        ),
    )
